"""
File parser service - handles PDF/DOCX parsing and Section 3 extraction.
"""
import io
import re
from pathlib import Path
from typing import Optional, Tuple

from pypdf import PdfReader
from docx import Document as DocxDocument


class FileParser:
    """Parse PDF and DOCX files, extract Section 3 content."""
    
    def __init__(self):
        # Section 3 heading patterns (case-insensitive)
        self.section_patterns = [
            r"^3\s*\.\s*trial\s+objectives\s+and\s+estimands?",
            r"^3\s*\.\s*objectives?\s+and\s+estimands?",
            r"^3\s*\.\s*study\s+objectives?\s+and\s+estimands?",
            r"^3\s+.*objectives.*estimand",
            r"^section\s+3\s*[:\.\-]?\s*.*objectives",
        ]
        self.section_regex = re.compile(
            "|".join(self.section_patterns), 
            re.MULTILINE | re.IGNORECASE
        )
        
        # Next section patterns to detect end of Section 3
        self.next_section_patterns = [
            r"^4\s*\.\s*trial\s+design",
            r"^4\s*\.\s*study\s+design",
            r"^section\s+4\s*[:\.\-]?\s*",
            r"^5\s*\.\s*selection\s+and\s+withdrawal",
            r"^section\s+5\s*[:\.\-]?\s*",
        ]
        self.next_section_regex = re.compile(
            "|".join(self.next_section_patterns),
            re.MULTILINE | re.IGNORECASE
        )
    
    def parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text_parts = []
        reader = PdfReader(str(file_path))
        
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        return "\n\n".join(text_parts)
    
    def parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(str(file_path))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    
    def parse_file(self, file_path: Path) -> str:
        """Parse file based on extension."""
        suffix = file_path.suffix.lower()
        
        if suffix == ".pdf":
            return self.parse_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    
    def extract_section_3(self, full_text: str) -> Tuple[str, bool]:
        """
        Extract Section 3 from full document text.
        
        Returns:
            Tuple of (section_3_text, was_found)
        """
        # Find Section 3 heading
        section_match = self.section_regex.search(full_text)
        
        if not section_match:
            return "", False
        
        # Get starting position
        start_pos = section_match.start()
        
        # Find next section to determine end
        remaining_text = full_text[start_pos:]
        next_section_match = self.next_section_regex.search(remaining_text)
        
        if next_section_match:
            # End at the next section heading
            end_pos = start_pos + next_section_match.start()
        else:
            # No clear next section, look for common endings
            # Try to find end of Section 3 content (usually a few thousand chars)
            end_pos = start_pos + 10000  # reasonable max
        
        section_3_text = full_text[start_pos:end_pos].strip()
        
        # Clean up the extracted text
        section_3_text = self._clean_text(section_3_text)
        
        return section_3_text, True
    
    def _clean_text(self, text: str) -> str:
        """Clean up extracted text."""
        # Remove excessive whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = re.sub(r" {2,}", " ", text)
        
        # Remove page numbers that might be in the text
        text = re.sub(r"\n\s*Page\s+\d+\s+of\s+\d+\s*\n", "\n", text, flags=re.IGNORECASE)
        
        return text.strip()
    
    def parse_and_extract(self, file_path: Path) -> Tuple[str, bool]:
        """
        Parse file and extract Section 3.
        
        Returns:
            Tuple of (section_3_text, was_found)
        """
        full_text = self.parse_file(file_path)
        return self.extract_section_3(full_text)


def get_parser() -> FileParser:
    """Get FileParser instance."""
    return FileParser()
